"""Markdown -> .docx converter tuned for VisibilityOS docs.

Handles: H1-H4, tables, fenced code blocks, bullet/numbered lists,
horizontal rules, inline bold/italic/code, links (text only).
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*\s][^*]*\*)")


def add_runs(paragraph, text: str) -> None:
    # Strip markdown links to their text
    text = LINK_RE.sub(r"\1", text)
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            shade(run)
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def shade(run) -> None:
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F1F3F5")
    rPr.append(shd)


def parse_table_block(lines: list[str], i: int) -> tuple[list[list[str]], int]:
    """Return (rows including header, new_index). Assumes lines[i] is the header row
    and lines[i+1] is the separator (|---|---|)."""
    header = [c.strip() for c in lines[i].strip().strip("|").split("|")]
    rows = [header]
    j = i + 2
    while j < len(lines):
        row = lines[j]
        if "|" not in row or not row.strip():
            break
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        # pad / trim to header width
        if len(cells) < len(header):
            cells += [""] * (len(header) - len(cells))
        elif len(cells) > len(header):
            cells = cells[: len(header)]
        rows.append(cells)
        j += 1
    return rows, j


def is_table_start(lines: list[str], i: int) -> bool:
    if i + 1 >= len(lines):
        return False
    if "|" not in lines[i]:
        return False
    sep = lines[i + 1].strip()
    if not sep:
        return False
    if "|" not in sep:
        return False
    # Separator must be made of `-`, `:`, `|`, and whitespace only
    return bool(re.match(r"^[\s\-:|]+$", sep))


def convert(md_path: Path, docx_path: Path) -> None:
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Tighten heading sizes a touch (default H1 is 16pt; we leave alone)
    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.startswith("```"):
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            shade(run)
            continue

        # Headings
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)

        # Table
        elif is_table_start(lines, i):
            rows, new_i = parse_table_block(lines, i)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        cell = table.rows[r_idx].cells[c_idx]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        if r_idx == 0:
                            run = p.add_run(LINK_RE.sub(r"\1", cell_text))
                            run.bold = True
                        else:
                            add_runs(p, cell_text)
            i = new_i
            continue

        # Bullet list
        elif re.match(r"^[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, re.sub(r"^[-*]\s+", "", line))

        # Numbered list
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\.\s+", "", line))

        # Horizontal rule
        elif line.strip() == "---":
            p = doc.add_paragraph()
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        # Blank line
        elif line.strip() == "":
            pass

        # Italic-only line (often used as a subtitle)
        elif re.match(r"^\*[^*].+\*$", line.strip()):
            p = doc.add_paragraph()
            run = p.add_run(line.strip()[1:-1])
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Plain paragraph
        else:
            p = doc.add_paragraph()
            add_runs(p, line)

        i += 1

    doc.save(docx_path)
    print(f"Wrote {docx_path} ({docx_path.stat().st_size:,} bytes)")


if __name__ == "__main__":
    here = Path(__file__).resolve().parent
    for name in ("VISIBILITY_OS_COMPLETE_REPORT", "VISIBILITY_OS_SUMMARY"):
        md = here / f"{name}.md"
        docx = here / f"{name}.docx"
        if md.exists():
            convert(md, docx)
        else:
            print(f"Skip: {md} not found")
